#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历生成器 — 严格复刻「WXX 简历模版.docx」的排版格式。

用法:
    python build_resume.py --content resume.json --output 我的简历.docx

内容 JSON 结构参见 resume.example.json 与 references/format-spec.md。
"""
import argparse
import json
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

# ---------------- 模板常量（严格来自模板 XML 解析） ----------------
FONT_CN = "宋体"
FONT_EN = "Times New Roman"      # 仅联系方式、求职意向的西文用此字体

SZ_NAME = 28        # 姓名 14pt（sz=28）
SZ_BODY = 21        # 正文/基本信息/时间 10.5pt（sz=21）
SZ_TITLE = 26       # 分区标题 13pt（sz=26）
SZ_ENTRY = 24       # 条目行 12pt（sz=24）

CHAR_SP_NAME = 4    # 姓名行、基本信息行字符间距
CHAR_SP_SUB = 9     # 联系方式、求职意向字符间距

TAB_RIGHT_POS = Cm(21 - 1.27 - 1.27)  # 右对齐制表位：可用宽度尽头 18.46cm（相对于左边距）
TAB_CENTER_POS = Cm((21 - 1.27 - 1.27) / 2)  # 居中制表位：可用宽度中间 9.23cm


# ---------------- 底层工具 ----------------
def _set_run_font(run, size_half=SZ_BODY, bold=False, cn=FONT_CN, en=FONT_CN,
                  char_spacing=None):
    """设置 run 的字体，使用 python-docx 标准 API（兼容性最好）。"""
    run.font.bold = bold
    run.font.name = en
    # 通过 XML 设置 eastAsia 字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:cs"), en)
    # Tencent Docs 只识别 w:val="1"/"0"，强制设置
    b = rpr.find(qn("w:b"))
    if b is not None:
        b.set(qn("w:val"), "1" if bold else "0")
    bcs = rpr.find(qn("w:bCs"))
    if bcs is None:
        bcs = OxmlElement("w:bCs")
        rpr.append(bcs)
    bcs.set(qn("w:val"), "1" if bold else "0")
    run.font.size = Pt(size_half / 2.0)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if char_spacing:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(char_spacing))
        rpr.append(sp)


def _set_para(para, align=None, before_twips=None, before_lines=None,
              after_twips=None, after_lines=None):
    """设置段落属性；before_lines/after_lines 以 1/100 行表示段前/段后。"""
    pf = para.paragraph_format
    if align is not None:
        para.alignment = align
    # 单倍行距（模板全文 line=240 lineRule=auto）
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # 段前/段后
    if before_twips is not None:
        pf.space_before = Twips(before_twips)
    if after_twips is not None:
        pf.space_after = Twips(after_twips)

    # 段前/段后按行数表达（w:beforeLines / w:afterLines）
    ppr = para._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    if before_lines is not None:
        spacing.set(qn("w:beforeLines"), str(before_lines))
        spacing.set(qn("w:beforeAutospacing"), "0")
    if after_lines is not None:
        spacing.set(qn("w:afterLines"), str(after_lines))
        spacing.set(qn("w:afterAutospacing"), "0")


def _add_bottom_border(para, sz=4, color="000000"):
    """段落底部加黑色单线（分区标题下划线）。"""
    ppr = para._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def _setup_section(doc):
    """页面设置：A4 + 上下 1cm + 左右 1.27cm + 文档网格。"""
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Twips(567)
    sec.bottom_margin = Twips(567)
    sec.left_margin = Twips(720)
    sec.right_margin = Twips(720)
    sec.header_distance = Twips(851)
    sec.footer_distance = Twips(992)
    # 单栏（与模板 w:cols num=1 space=0 一致）
    sect_pr = sec._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "1")
    cols.set(qn("w:space"), "0")
    # 文档网格 lines / linePitch=312（先查找已有的，避免重复）
    doc_grid = sect_pr.find(qn("w:docGrid"))
    if doc_grid is None:
        doc_grid = OxmlElement("w:docGrid")
        sect_pr.append(doc_grid)
    doc_grid.set(qn("w:type"), "lines")
    doc_grid.set(qn("w:linePitch"), "312")
    doc_grid.set(qn("w:charSpace"), "0")


# ---------------- 段落构建（严格按模板） ----------------
def _add_para(doc, runs, align=WD_ALIGN_PARAGRAPH.LEFT, **ppr_kwargs):
    """添加一段，runs 为 (text, fmt_kwargs) 列表。
    同时把第一个 run 的加粗复制到段落级 rPr（Tencent Docs 需要段落级加粗属性）。"""
    para = doc.add_paragraph()
    _set_para(para, align=align, **ppr_kwargs)
    for text, fmt in runs:
        run = para.add_run(text)
        _set_run_font(run, **fmt)
    # 段落级 rPr：复制第一个 run 的加粗属性
    if runs:
        bold_val = runs[0][1].get("bold", False)
        ppr = para._p.get_or_add_pPr()
        para_rpr = ppr.find(qn("w:rPr"))
        if para_rpr is None:
            para_rpr = OxmlElement("w:rPr")
            ppr.append(para_rpr)
        b = para_rpr.find(qn("w:b"))
        if b is None:
            b = OxmlElement("w:b")
            para_rpr.append(b)
        b.set(qn("w:val"), "1" if bold_val else "0")
    return para


def add_name(doc, name):
    """姓名行：居中，14pt 加粗，字符间距4，段前42twips。
    模板 P1: ascii=宋体 hAnsi=宋体 sz=28 bold=True char_spacing=4 before=42 center"""
    _add_para(doc,
              [(name, dict(size_half=SZ_NAME, bold=True, cn=FONT_CN, en=FONT_CN,
                           char_spacing=CHAR_SP_NAME))],
              align=WD_ALIGN_PARAGRAPH.CENTER,
              before_twips=42)


def add_basic_info(doc, text):
    """基本信息行：居中，10.5pt 不加粗，全宋体，字符间距4，段前44twips。
    模板 P2: ascii=宋体 eastAsia=宋体 hAnsi=宋体 sz=21 bold=False char_spacing=4 before=44 center"""
    _add_para(doc,
              [(text, dict(size_half=SZ_BODY, bold=False, cn=FONT_CN, en=FONT_CN,
                           char_spacing=CHAR_SP_NAME))],
              align=WD_ALIGN_PARAGRAPH.CENTER,
              before_twips=44)


def add_contact(doc, text):
    """联系方式行：居中，10.5pt 不加粗，西文Times New Roman 中文宋体，字符间距9，段前44twips。
    模板 P3: ascii=Times New Roman eastAsia=宋体 sz=21 char_spacing=9 before=44 center"""
    _add_para(doc,
              [(text, dict(size_half=SZ_BODY, bold=False, cn=FONT_CN, en=FONT_EN,
                           char_spacing=CHAR_SP_SUB))],
              align=WD_ALIGN_PARAGRAPH.CENTER,
              before_twips=44)


def add_job_intent(doc, text):
    """求职意向行：居中，10.5pt 不加粗，西文Times New Roman 中文宋体，字符间距9，段前44twips。
    用户要求：求职意向在手机号下面，居中。
    模板 P4: ascii=Times New Roman eastAsia=宋体 sz=21 char_spacing=9 before=44
    （模板XML为both+空格，但用户明确要求居中，视觉效果一致）"""
    _add_para(doc,
              [(text, dict(size_half=SZ_BODY, bold=False, cn=FONT_CN, en=FONT_EN,
                           char_spacing=CHAR_SP_SUB))],
              align=WD_ALIGN_PARAGRAPH.CENTER,
              before_twips=44)


def add_blank_line(doc):
    """求职意向后的空段落（匹配模板 P5）。
    模板 P5: before=44 ind firstLine=2178 firstLineChars=1100 char_spacing=9"""
    para = doc.add_paragraph()
    _set_para(para, before_twips=44)
    ppr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "2178")
    ind.set(qn("w:firstLineChars"), "1100")
    ppr.append(ind)


def add_section_title(doc, title, is_first=False):
    """分区标题：13pt 加粗，底部黑色下划线，段前段后各0.5行。
    模板 P6(首个): before=0 after=159 afterLines=50
    模板 P12(后续): before=159 beforeLines=50 after=159 afterLines=50"""
    if is_first:
        # 首个标题段前=0（前面已有空段落间隔）
        para = _add_para(doc,
                         [(title, dict(size_half=SZ_TITLE, bold=True,
                                       cn=FONT_CN, en=FONT_CN))],
                         align=WD_ALIGN_PARAGRAPH.LEFT,
                         after_twips=159, after_lines=50)
    else:
        para = _add_para(doc,
                         [(title, dict(size_half=SZ_TITLE, bold=True,
                                       cn=FONT_CN, en=FONT_CN))],
                         align=WD_ALIGN_PARAGRAPH.LEFT,
                         before_twips=159, before_lines=50,
                         after_twips=159, after_lines=50)
    _add_bottom_border(para, sz=4, color="000000")


def add_entry(doc, left_text, center_text="", right_text=""):
    """条目行：左侧名称（加粗12pt）+ 居中角色名（加粗12pt）+ 右对齐时间（不加粗10.5pt）。
    用户要求：项目负责人/人力资源部干事等角色名要在这一排居中。
    全宋体。时间永远不加粗。"""
    para = doc.add_paragraph()
    _set_para(para, align=WD_ALIGN_PARAGRAPH.LEFT)
    # 居中制表位（页面中间）+ 右对齐制表位（右边距处）
    para.paragraph_format.tab_stops.add_tab_stop(TAB_CENTER_POS, WD_TAB_ALIGNMENT.CENTER)
    para.paragraph_format.tab_stops.add_tab_stop(TAB_RIGHT_POS, WD_TAB_ALIGNMENT.RIGHT)
    # 左侧名称
    run1 = para.add_run(left_text)
    _set_run_font(run1, size_half=SZ_ENTRY, bold=True, cn=FONT_CN, en=FONT_CN)
    if center_text:
        # 居中角色名
        run2 = para.add_run("\t" + center_text)
        _set_run_font(run2, size_half=SZ_ENTRY, bold=True, cn=FONT_CN, en=FONT_CN)
    if right_text:
        # 右对齐时间：如果没有 center_text，需要两个 \t 跳过居中制表位
        tabs = "\t\t" if not center_text else "\t"
        run3 = para.add_run(tabs + right_text)
        _set_run_font(run3, size_half=SZ_BODY, bold=False, cn=FONT_CN, en=FONT_CN)
    return para


def add_body(doc, text):
    """普通正文段：10.5pt 不加粗 全宋体。
    模板 P8/P9/P10: ascii=宋体 eastAsia=宋体 sz=21 bold=False"""
    _add_para(doc,
              [(text, dict(size_half=SZ_BODY, bold=False, cn=FONT_CN, en=FONT_CN))],
              align=WD_ALIGN_PARAGRAPH.LEFT)


def add_label_body(doc, label, text=""):
    """标签（恒加粗）+ 正文（恒不加粗）。全宋体。
    模板 P17: run0(标签) bold=True; run1(正文) bold=False
    模板 P18: run0(标签) bold=True; run1(正文) bold=False
    标签包括：项目背景：/技术栈：/代码仓库：/主要职责：/项目成果：/核心工作：/项目描述："""
    runs = [(label, dict(size_half=SZ_BODY, bold=True, cn=FONT_CN, en=FONT_CN))]
    if text:
        runs.append((text, dict(size_half=SZ_BODY, bold=False, cn=FONT_CN, en=FONT_CN)))
    _add_para(doc, runs, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_bullet(doc, text):
    """要点行：• 前缀 + 悬挂缩进，10.5pt 不加粗 全宋体。
    用户要求：主要职责/项目成果后面的内容要分点，每行一个•标记。"""
    para = _add_para(doc,
                     [("\u2022 " + text, dict(size_half=SZ_BODY, bold=False,
                                  cn=FONT_CN, en=FONT_CN))],
                     align=WD_ALIGN_PARAGRAPH.LEFT)
    # 悬挂缩进：left=420 hanging=420 firstLineChars=0（让第二行与•后的文字对齐）
    ppr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "420")
    ind.set(qn("w:leftChars"), "0")
    ind.set(qn("w:hanging"), "420")
    ind.set(qn("w:firstLineChars"), "0")
    ppr.append(ind)
    return para


# ---------------- 主流程 ----------------
def build(content, output):
    doc = Document()
    _setup_section(doc)

    # 设置 Normal 样式：消除 python-docx 默认的 space_after=8pt 和 line_spacing=1.08
    # 并添加模板 Normal 样式特有的属性（snapToGrid=0 等，解决行间距宽的问题）
    normal = doc.styles['Normal']
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.font.name = FONT_CN
    normal.font.size = Pt(SZ_BODY / 2.0)
    # Normal 样式的 eastAsia 字体设为宋体
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_CN)
    rfonts.set(qn("w:ascii"), FONT_CN)
    rfonts.set(qn("w:hAnsi"), FONT_CN)
    # kern=0: 禁用字距调整
    kern = OxmlElement("w:kern")
    kern.set(qn("w:val"), "0")
    rpr.append(kern)
    # 段落级属性（与模板 Normal 样式一致）
    ppr = normal.element.get_or_add_pPr()
    for tag in ("w:snapToGrid", "w:autoSpaceDE", "w:autoSpaceDN",
                "w:adjustRightInd", "w:kinsoku"):
        elem = OxmlElement(tag)
        elem.set(qn("w:val"), "0")
        ppr.append(elem)

    # 页首个人信息区（顺序固定）
    add_name(doc, content["name"])
    add_basic_info(doc, content["basic_info"])
    add_contact(doc, content["contact"])
    add_job_intent(doc, content["job_intent"])
    add_blank_line(doc)  # 求职意向后的空段落

    # 分区
    for i, section in enumerate(content["sections"]):
        add_section_title(doc, section["title"], is_first=(i == 0))
        for item in section.get("items", []):
            itype = item.get("type", "body")
            if itype == "entry":
                add_entry(doc, item["left"], item.get("center", ""),
                          item.get("right", ""))
            elif itype == "body":
                add_body(doc, item["text"])
            elif itype == "label_body":
                add_label_body(doc, item["label"], item.get("text", ""))
            elif itype == "bullet":
                add_bullet(doc, item["text"])
            else:
                raise ValueError(f"未知 item 类型: {itype}")

    doc.save(output)
    print(f"已生成: {output}")


def main():
    ap = argparse.ArgumentParser(description="按 WXX 简历模版格式生成简历 docx")
    ap.add_argument("--content", required=True, help="内容 JSON 文件路径")
    ap.add_argument("--output", required=True, help="输出 docx 路径")
    args = ap.parse_args()

    with open(args.content, "r", encoding="utf-8") as f:
        content = json.load(f)

    build(content, args.output)


if __name__ == "__main__":
    main()
